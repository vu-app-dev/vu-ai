import io
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from models.cv import CvAnalyzeResponse, CvDimensions
from services.cv.cv_analyzer import CvAnalyzer


class TestGetExtension:
    def test_pdf_extension(self):
        result = CvAnalyzer()._get_extension("https://example.com/resume.pdf")
        assert result == ".pdf"

    def test_docx_extension(self):
        result = CvAnalyzer()._get_extension("https://example.com/resume.docx")
        assert result == ".docx"

    def test_unsupported_extension(self):
        result = CvAnalyzer()._get_extension("https://example.com/resume.txt")
        assert result == ""

    def test_extension_with_query_params(self):
        result = CvAnalyzer()._get_extension("https://example.com/resume.pdf?token=abc")
        assert result == ".pdf"

    def test_extension_with_fragment(self):
        result = CvAnalyzer()._get_extension("https://example.com/resume.docx#page=1")
        assert result == ".docx"


class TestExtractPdf:
    def test_extract_pdf_handles_invalid_content(self):
        result = CvAnalyzer()._extract_text(b"not a real pdf", ".pdf")
        assert result is None

    def test_extract_docx_text(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("John Doe - Software Engineer")
        doc.add_paragraph("Skills: Python, React, Node.js")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = CvAnalyzer._extract_docx(content)
        assert "John Doe" in result
        assert "Python" in result


class TestAnalyzeUnsupported:
    @pytest.mark.asyncio
    async def test_unsupported_file_type_returns_none(self):
        analyzer = CvAnalyzer()
        result = await analyzer.analyze("https://example.com/resume.txt")
        assert result is None

    @pytest.mark.asyncio
    async def test_empty_url_extension(self):
        analyzer = CvAnalyzer()
        result = await analyzer.analyze("https://example.com/file")
        assert result is None


class TestCvDimensionsScoring:
    def test_compute_score_with_dimensions(self):
        response = CvAnalyzeResponse(
            skills=["Python", "React"],
            summary="Test summary",
            dimensions=CvDimensions(
                skillsMatch=4, experienceDepth=3, educationFit=3, projectRelevance=5
            ),
        )
        score = response.compute_score()
        assert score == 76.0

    def test_compute_score_all_fives(self):
        response = CvAnalyzeResponse(
            skills=[],
            summary="",
            dimensions=CvDimensions(
                skillsMatch=5, experienceDepth=5, educationFit=5, projectRelevance=5
            ),
        )
        assert response.compute_score() == 100.0

    def test_compute_score_all_ones(self):
        response = CvAnalyzeResponse(
            skills=[],
            summary="",
            dimensions=CvDimensions(
                skillsMatch=1, experienceDepth=1, educationFit=1, projectRelevance=1
            ),
        )
        assert response.compute_score() == 20.0

    def test_compute_score_no_dimensions(self):
        response = CvAnalyzeResponse(skills=[], summary="")
        assert response.compute_score() is None


class TestAnalyzeWithLLM:
    @pytest.mark.asyncio
    async def test_analyze_cv_llm_success(self):
        dimensions = CvDimensions(
            skillsMatch=4, experienceDepth=4, educationFit=3, projectRelevance=5
        )
        mock_llm = AsyncMock()
        mock_llm.generate_json = AsyncMock(return_value=CvAnalyzeResponse(
            skills=["Python", "React", "Node.js"],
            summary="Experienced full-stack developer with 5 years of experience.",
            dimensions=dimensions,
        ))
        analyzer = CvAnalyzer(llm=mock_llm)
        mock_content = b"dummy pdf content"

        with patch.object(analyzer, "_download", new_callable=AsyncMock) as mock_dl, \
             patch.object(analyzer, "_extract_text", return_value="John Doe, Software Engineer"):
            mock_dl.return_value = mock_content
            result = await analyzer.analyze(
                "https://example.com/resume.pdf",
                job_context={"role": "Senior Developer"},
            )
            assert result is not None
            assert "Python" in result.skills
            assert result.dimensions is not None
            assert result.score is not None
            assert result.score == result.compute_score()

    @pytest.mark.asyncio
    async def test_analyze_cv_llm_failure_returns_null_scores(self):
        mock_llm = AsyncMock()
        mock_llm.generate_json = AsyncMock(side_effect=Exception("LLM error"))
        analyzer = CvAnalyzer(llm=mock_llm)

        with patch.object(analyzer, "_download", new_callable=AsyncMock) as mock_dl, \
             patch.object(analyzer, "_extract_text", return_value="John Doe, Software Engineer"):
            mock_dl.return_value = b"dummy content"
            result = await analyzer.analyze("https://example.com/resume.pdf")
            assert result is not None
            assert result.score is None
            assert result.dimensions is None

    @pytest.mark.asyncio
    async def test_analyze_cv_download_failure(self):
        analyzer = CvAnalyzer()
        with patch.object(analyzer, "_download", new_callable=AsyncMock, return_value=None):
            result = await analyzer.analyze("https://unreachable.example.com/resume.pdf")
            assert result is None

    @pytest.mark.asyncio
    async def test_analyze_cv_empty_extraction(self):
        mock_llm = AsyncMock()
        analyzer = CvAnalyzer(llm=mock_llm)

        with patch.object(analyzer, "_download", new_callable=AsyncMock, return_value=b"content"), \
             patch.object(analyzer, "_extract_text", return_value="   "):
            result = await analyzer.analyze("https://example.com/resume.pdf")
            assert result is not None
            assert result.skills == []
            assert result.dimensions is None